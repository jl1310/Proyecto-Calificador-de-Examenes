from docx import Document
import os

def create_report(data, output_path="reporte.docx"):
    doc = Document()

    # Título del documento
    doc.add_heading("Informe de Calificación de Examen", 0)

    # Información general
    doc.add_paragraph(f"📌 ID del Examen: {data.get('examen_id', 'Desconocido')}")
    doc.add_paragraph(f"👤 Estudiante: {data.get('nombre_estudiante', 'No identificado')}")
    doc.add_paragraph(f"📊 Puntaje Total: {data.get('puntaje_total', 0)}\n")

    # Resultados por pregunta
    doc.add_heading("Resultados por Pregunta", level=1)

    for pregunta in data.get("preguntas", []):
        doc.add_heading(f"Pregunta {pregunta.get('numero', '?')}", level=2)
        doc.add_paragraph(f"📝 Respuesta del estudiante:\n{pregunta.get('respuesta', '')}")
        doc.add_paragraph(f"🎯 Puntaje: {pregunta.get('puntaje', 0)} / 2")
        doc.add_paragraph(f"📌 Justificación: {pregunta.get('justificacion', 'Sin justificación.')}")
        doc.add_paragraph("")  # espacio adicional

    # Conclusión general
    doc.add_heading("🧠 Conclusión General del Desempeño", level=1)

    conclusion = data.get("conclusion", {})
    doc.add_paragraph(f"✅ Fortalezas: {conclusion.get('fortalezas', 'No especificadas')}")
    doc.add_paragraph(f"⚠️ Debilidades: {conclusion.get('debilidades', 'No especificadas')}")
    doc.add_paragraph(f"📚 Sugerencias de Mejora: {conclusion.get('sugerencias', 'No especificadas')}")

    # Crear carpeta si no existe
    if not os.path.exists("reportes"):
        os.makedirs("reportes")

    file_path = os.path.join("reportes", output_path)
    doc.save(file_path)

    return file_path